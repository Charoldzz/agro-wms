from __future__ import annotations

import argparse
import io
import json
import shutil
import subprocess
from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.colors import black, white
from reportlab.pdfgen import canvas


MAX_TEMPLATE_ROWS = 8


def fmt(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def replace_in_paragraph(paragraph, values: dict[str, str]) -> None:
    original = paragraph.text
    text = original
    for key, value in values.items():
        text = text.replace("{{" + key + "}}", fmt(value))
    if text == original:
        return

    # Word often splits placeholders across runs. Rebuild the paragraph text
    # when a replacement is needed so every marker is filled consistently.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def replace_in_table(table, values: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, values)


def build_values(payload: dict) -> dict[str, str]:
    items = payload.get("items") or []
    values = {
        "Number": payload.get("number", ""),
        "Fecha": payload.get("date", ""),
        "Empresa": payload.get("client", ""),
        "Trans": payload.get("driver_name") or payload.get("receiver_name") or "",
        "Contacto": payload.get("contact") or payload.get("receiver_name") or "",
        "Placa": payload.get("vehicle_plate", ""),
        "Observaciones": payload.get("notes", ""),
        "Recibido": payload.get("received_by", ""),
        "Entregado": payload.get("delivered_by") or payload.get("user_email", ""),
    }

    for index in range(1, MAX_TEMPLATE_ROWS + 1):
        item = items[index - 1] if index <= len(items) else {}
        quantity = item.get("quantity", "")
        package_size = float(item.get("package_size") or 0)
        package_unit = item.get("package_unit") or ""
        equivalent = ""
        if quantity != "" and package_size:
            equivalent = f"{float(quantity) * package_size:g} {package_unit}".strip()
        values[f"Cantidad{index}"] = quantity
        values[f"Volumen{index}"] = equivalent
        values[f"Producto{index}"] = item.get("product", "")
        values[f"CantE{index}"] = (
            item.get("box_count")
            or item.get("packaging")
            or item.get("quantity")
            or ""
        )

    return values


def fill_template(template_path: Path, payload: dict, output_docx: Path) -> None:
    doc = Document(template_path)
    values = build_values(payload)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, values)
    for table in doc.tables:
        replace_in_table(table, values)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "No se encontro LibreOffice/soffice. Instala LibreOffice en el servidor "
            "o usa este script solo para generar DOCX."
        )
    output_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return output_dir / f"{docx_path.stem}.pdf"


def draw_clean_text(c, x: float, y: float, w: float, h: float, text: str, size: int = 9, bold: bool = False) -> None:
    c.setFillColor(white)
    c.rect(x, y, w, h, stroke=0, fill=1)
    c.setFillColor(black)
    c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
    c.drawString(x + 2, y + 4, fmt(text)[:120])


def fill_pdf_template(template_path: Path, payload: dict, output_pdf: Path) -> None:
    reader = PdfReader(str(template_path))
    page = reader.pages[0]
    width = float(page.mediabox.width)
    height = float(page.mediabox.height)

    packet = io.BytesIO()
    c = canvas.Canvas(packet, pagesize=(width, height))
    values = build_values(payload)

    draw_clean_text(c, 682, 540, 105, 24, values["Number"], 12, True)
    draw_clean_text(c, 704, 512, 82, 22, values["Fecha"], 10)
    draw_clean_text(c, 120, 486, 280, 22, values["Empresa"], 10, True)
    draw_clean_text(c, 510, 486, 210, 22, values["Trans"], 10)
    draw_clean_text(c, 145, 459, 255, 22, values["Contacto"], 10)
    draw_clean_text(c, 520, 459, 170, 22, values["Placa"], 10, True)

    row_ys = [394, 378, 362, 346, 330, 314, 298, 282]
    for index, y in enumerate(row_ys, start=1):
        draw_clean_text(c, 50, y, 78, 15, values[f"Cantidad{index}"], 8)
        draw_clean_text(c, 137, y, 86, 15, values[f"Volumen{index}"], 8)
        draw_clean_text(c, 225, y, 230, 15, values[f"Producto{index}"], 7)
        draw_clean_text(c, 457, y, 50, 15, values[f"CantE{index}"], 8)
        draw_clean_text(c, 510, y, 50, 15, "", 8)
        draw_clean_text(c, 570, y, 52, 15, "", 8)
        draw_clean_text(c, 632, y, 58, 15, "", 8)
        draw_clean_text(c, 700, y, 50, 15, "", 8)

    # The observation placeholder is outside the table. Leave the title intact and
    # replace only the marker area with wrapped plain text.
    draw_clean_text(c, 34, 198, 720, 32, "", 8)
    c.setFillColor(black)
    c.setFont("Helvetica", 9)
    notes = fmt(values["Observaciones"])
    for i, line in enumerate([notes[j:j + 120] for j in range(0, len(notes), 120)][:3]):
        c.drawString(38, 218 - (i * 11), line)

    draw_clean_text(c, 252, 32, 150, 22, values["Recibido"], 10)
    draw_clean_text(c, 535, 32, 190, 22, values["Entregado"], 10)
    c.save()

    packet.seek(0)
    overlay = PdfReader(packet)
    page.merge_page(overlay.pages[0])

    writer = PdfWriter()
    writer.add_page(page)
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    with output_pdf.open("wb") as fh:
        writer.write(fh)


def main() -> None:
    parser = argparse.ArgumentParser(description="Rellena una orden de almacen en DOCX o PDF.")
    parser.add_argument("--type", choices=["ingreso", "salida"], required=True)
    parser.add_argument("--payload", required=True, help="Archivo JSON con los datos de la operacion.")
    parser.add_argument("--out", required=True, help="Ruta de salida DOCX o PDF.")
    parser.add_argument("--format", choices=["pdf", "docx"], default="pdf")
    parser.add_argument("--pdf", action="store_true", help="Convierte el DOCX a PDF usando LibreOffice/soffice.")
    args = parser.parse_args()

    root = Path(__file__).resolve().parents[1]
    payload = json.loads(Path(args.payload).read_text(encoding="utf-8-sig"))
    output_path = Path(args.out)

    if args.format == "pdf":
        template = root / "templates" / "warehouse" / (
            "orden_ingreso.pdf" if args.type == "ingreso" else "orden_salida.pdf"
        )
        fill_pdf_template(template, payload, output_path)
        print(output_path)
        return

    template = root / "templates" / "warehouse" / (
        "orden_ingreso.docx" if args.type == "ingreso" else "orden_salida.docx"
    )
    fill_template(template, payload, output_path)

    if args.pdf:
        pdf_path = convert_to_pdf(output_path, output_path.parent)
        print(pdf_path)
    else:
        print(output_path)


if __name__ == "__main__":
    main()
